import docx
from docx.shared import Inches #添加图片的包

'''
1、Document代表整个word文件
2、Paragraphs 一个段落代表的是一个paragraphs对象
3、Run 代表相同字样的一段
'''
wdoc = docx.Document('练习文件/刘亦菲简介.docx')
print('段落数总共 =',len(wdoc.paragraphs))
for i in range(0,len(wdoc.paragraphs)):
    print('paragraphs =%d'%i,wdoc.paragraphs[i].text)

print('paragraphs 1的Run数量 = ',len(wdoc.paragraphs[1].runs))
for i in range(0,len(wdoc.paragraphs[1].runs)):
    print('Run =%d'%i,wdoc.paragraphs[1].runs[i].text)   

print('paragraphs 1的Run数量 = ',len(wdoc.paragraphs[2].runs))
for i in range(0,len(wdoc.paragraphs[2].runs)):
    print('Run =%d'%i,wdoc.paragraphs[2].runs[i].text) 

#save()文件进行复制
wdoc.save('练习文件/刘亦菲简介副本.docx') 
#首先设置标题，标题格式，最后保存在docx里，打开就可以看到设置的格式
wdoc1 = docx.Document()
wdoc1.add_heading('刘亦菲',level=3)
prt=wdoc1.add_paragraph('我是第一个段落')
wdoc1.add_paragraph('我是第二个段落')
wdoc1.add_paragraph('刘亦菲照片')
#指定图片进行插入，并设置大小
wdoc1.add_picture('/Users/wangrui/Desktop/Python练习汇总文档/练习文件/IMG_3869.JPG',width=Inches(3.0))
wdoc1.add_picture('/Users/wangrui/Desktop/Python练习汇总文档/练习文件/IMG_3870.JPG',width=Inches(3.0))
prt.insert_paragraph_before('我是插队进来的') #将文本添加到首行
wdoc1.add_page_break() #强制换页,在新的一页加文字
wdoc1.add_paragraph('我是新的一个段落')
print(wdoc1.paragraphs[0].text)
wdoc1.save('练习.docx')

#建立表格
'''
rows代表的是行
cols代表的是列
'''
wdoc2=docx.Document()
table=wdoc2.add_table(rows=3,cols=2)
row=table.rows[0]
row.cells[0].text = '姓名'
row.cells[1].text = '年龄'
row=table.rows[1]
row.cells[0].text = '刘安珂'
row.cells[1].text = '27'
row=table.rows[2]
row.cells[0].text = '金晶'
row.cells[1].text = '28'
#增加表格的列和输入数据
new_row = table.add_row()
new_row.cells[0].text = '夏千姿'
new_row.cells[1].text = '未知'
table.style = 'LightShading-Accent1'
wdoc2.save('人名统计.docx')

#段落样式
wdoc3=docx.Document()
wdoc3.add_paragraph('北京大学',style='ListNumber')
wdoc3.add_paragraph('清华大学',style='ListNumber')
wdoc3.add_paragraph('长治学院',style='ListNumber')
wdoc3.save('学校名称.docx')

#Run格式
wdoc4=docx.Document()
prt1=wdoc4.add_paragraph('我就是这么拽')
run1 = prt1.add_run('拽有错吗？')
run1.bold = True
run2 = prt1.add_run('你不止拽，还很得瑟。')
run2.strike = True
wdoc4.save('拽的定义.docx')




